# -*- coding: utf-8 -*-
"""Format exporters for BFSU ProofLens.

These functions are intentionally GUI-independent.
"""
from __future__ import annotations

import html
import json
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable
from xml.dom import minidom
from xml.etree import ElementTree as ET

from .utils import APP_NAME, APP_VERSION

DEFAULT_OPTIONS = {
    "include_ocr_text": True,
    "include_corrected_text": True,
    "include_final_text": True,
    "include_suggestions": True,
    "include_ocr_blocks": True,
    "final_text_only": False,
    "split_by_source_file": False,
    "current_file_id": "",
    "current_page_index": 0,
}


def _opts(options: dict | None) -> dict:
    merged = dict(DEFAULT_OPTIONS)
    merged.update(options or {})
    return merged


def _project_meta(project_data: dict[str, Any]) -> dict[str, Any]:
    settings = project_data.get("settings", {})
    return {
        "software": project_data.get("software", APP_NAME),
        "version": project_data.get("version", APP_VERSION),
        "project_name": project_data.get("project_name", "Untitled"),
        "export_time": datetime.now().isoformat(timespec="seconds"),
        "ocr_backend": settings.get("ocr", {}).get("backend", "rapidocr"),
        "use_llm": bool(settings.get("llm", {}).get("enabled", False)),
        "llm_model": settings.get("llm", {}).get("model", ""),
    }


def _iter_pages(project_data: dict[str, Any], scope: str = "project", options: dict | None = None):
    options = _opts(options)
    current_file_id = options.get("current_file_id", "")
    current_page_index = int(options.get("current_page_index") or 0)
    for file_entry in project_data.get("files", []):
        if scope in {"current_file", "file", "current_page", "page"} and current_file_id and file_entry.get("id") != current_file_id:
            continue
        pages = file_entry.get("pages", [])
        for idx, page in enumerate(pages):
            if scope in {"current_page", "page"} and idx != current_page_index:
                continue
            yield file_entry, idx, page


def _subset_project(project_data: dict[str, Any], scope: str, options: dict | None = None) -> dict[str, Any]:
    meta = deepcopy(project_data)
    files = []
    grouped: dict[str, dict[str, Any]] = {}
    for file_entry, _, page in _iter_pages(project_data, scope, options):
        fid = file_entry.get("id")
        if fid not in grouped:
            copied = {k: deepcopy(v) for k, v in file_entry.items() if k != "pages"}
            copied["pages"] = []
            grouped[fid] = copied
            files.append(copied)
        grouped[fid]["pages"].append(deepcopy(page))
    meta["files"] = files
    return meta


def _text_for_page(page: dict[str, Any]) -> str:
    return page.get("final_text") or page.get("corrected_text") or page.get("ocr_text") or ""


def _md_escape_cell(text: Any) -> str:
    s = str(text if text is not None else "")
    return s.replace("|", "\\|").replace("\n", "<br>")


def _bbox_str(block: dict[str, Any]) -> str:
    bbox = block.get("bbox") or []
    return json.dumps(bbox, ensure_ascii=False)


def export_to_txt(project_data: dict[str, Any], export_path: str | Path, scope: str = "project", options: dict | None = None) -> str:
    options = _opts(options)
    export_path = Path(export_path)
    export_path.parent.mkdir(parents=True, exist_ok=True)
    lines: list[str] = []
    if options.get("final_text_only"):
        for file_entry, _, page in _iter_pages(project_data, scope, options):
            lines.append(f"===== File: {file_entry.get('file_name', '')} | Page: {page.get('page_no', '')} =====")
            lines.append(_text_for_page(page))
            lines.append("")
    else:
        for file_entry, _, page in _iter_pages(project_data, scope, options):
            lines.append(f"===== File: {file_entry.get('file_name', '')} | Page: {page.get('page_no', '')} =====")
            if options.get("include_ocr_text"):
                lines.extend(["", "[OCR Text]", page.get("ocr_text", "")])
            if options.get("include_corrected_text"):
                lines.extend(["", "[Corrected Text]", page.get("corrected_text", "")])
            if options.get("include_final_text"):
                lines.extend(["", "[Final Text]", page.get("final_text", "")])
            if options.get("include_suggestions"):
                lines.append("\n[Suggestions]")
                for s in page.get("suggestions", []):
                    lines.append(f"- line {s.get('line_no', '')}: {s.get('original', '')} -> {s.get('suggested', '')} ({s.get('category', '')}, {s.get('confidence', '')}) {s.get('reason', '')}")
            lines.append("")
    export_path.write_text("\n".join(lines), encoding="utf-8")
    return str(export_path)


def export_to_json(project_data: dict[str, Any], export_path: str | Path, scope: str = "project", options: dict | None = None) -> str:
    export_path = Path(export_path)
    export_path.parent.mkdir(parents=True, exist_ok=True)
    subset = _subset_project(project_data, scope, options)
    meta = _project_meta(project_data)
    subset["export_time"] = meta["export_time"]
    export_path.write_text(json.dumps(subset, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(export_path)


def _add_text(parent: ET.Element, tag: str, value: Any) -> ET.Element:
    elem = ET.SubElement(parent, tag)
    elem.text = "" if value is None else str(value)
    return elem


def export_to_xml(project_data: dict[str, Any], export_path: str | Path, scope: str = "project", options: dict | None = None) -> str:
    options = _opts(options)
    meta = _project_meta(project_data)
    root = ET.Element("BFSUProofLensProject", attrib={"software": meta["software"], "version": meta["version"]})
    info = ET.SubElement(root, "ProjectInfo")
    _add_text(info, "ProjectName", meta["project_name"])
    _add_text(info, "ExportTime", meta["export_time"])
    _add_text(info, "OCRBackend", meta["ocr_backend"])
    _add_text(info, "UseLLM", str(meta["use_llm"]))
    _add_text(info, "LLMModel", meta["llm_model"])
    files_elem = ET.SubElement(root, "Files")
    file_nodes: dict[str, ET.Element] = {}
    for file_entry, _, page in _iter_pages(project_data, scope, options):
        fid = file_entry.get("id", file_entry.get("file_name", ""))
        if fid not in file_nodes:
            file_nodes[fid] = ET.SubElement(files_elem, "File", attrib={"name": str(file_entry.get("file_name", "")), "path": str(file_entry.get("file_path", ""))})
        page_elem = ET.SubElement(file_nodes[fid], "Page", attrib={"number": str(page.get("page_no", ""))})
        if options.get("include_ocr_text"):
            _add_text(page_elem, "OCRText", page.get("ocr_text", ""))
        if options.get("include_corrected_text"):
            _add_text(page_elem, "CorrectedText", page.get("corrected_text", ""))
        if options.get("include_final_text"):
            _add_text(page_elem, "FinalText", page.get("final_text", ""))
        if options.get("include_ocr_blocks"):
            blocks_elem = ET.SubElement(page_elem, "OCRBlocks")
            for i, block in enumerate(page.get("ocr_blocks", []), start=1):
                block_elem = ET.SubElement(blocks_elem, "Block", attrib={"id": str(i), "confidence": str(block.get("confidence", ""))})
                _add_text(block_elem, "Text", block.get("text", ""))
                bbox_elem = ET.SubElement(block_elem, "BBox")
                for point in block.get("bbox", []) or []:
                    x = point[0] if len(point) > 0 else ""
                    y = point[1] if len(point) > 1 else ""
                    ET.SubElement(bbox_elem, "Point", attrib={"x": str(x), "y": str(y)})
        if options.get("include_suggestions"):
            suggestions_elem = ET.SubElement(page_elem, "Suggestions")
            for s in page.get("suggestions", []):
                s_elem = ET.SubElement(suggestions_elem, "Suggestion", attrib={
                    "line_no": str(s.get("line_no", "")),
                    "category": str(s.get("category", "")),
                    "confidence": str(s.get("confidence", "")),
                    "status": str(s.get("status", "pending")),
                })
                _add_text(s_elem, "Original", s.get("original", ""))
                _add_text(s_elem, "Suggested", s.get("suggested", ""))
                _add_text(s_elem, "Reason", s.get("reason", ""))
    xml_bytes = ET.tostring(root, encoding="utf-8")
    pretty = minidom.parseString(xml_bytes).toprettyxml(indent="  ", encoding="utf-8")
    export_path = Path(export_path)
    export_path.parent.mkdir(parents=True, exist_ok=True)
    export_path.write_bytes(pretty)
    return str(export_path)


def export_to_markdown(project_data: dict[str, Any], export_path: str | Path, scope: str = "project", options: dict | None = None) -> str:
    options = _opts(options)
    meta = _project_meta(project_data)
    lines: list[str] = ["# BFSU ProofLens OCR 校对结果", "", "## 基本信息", ""]
    lines.extend([
        f"- 软件：{meta['software']}",
        f"- 版本：{meta['version']}",
        f"- 项目名称：{meta['project_name']}",
        f"- 导出时间：{meta['export_time']}",
        f"- OCR 引擎：{meta['ocr_backend']}",
        f"- 是否使用 LLM：{meta['use_llm']}",
        f"- LLM 模型：{meta['llm_model']}",
        "",
    ])
    current_file_name = None
    for file_entry, _, page in _iter_pages(project_data, scope, options):
        if file_entry.get("file_name") != current_file_name:
            current_file_name = file_entry.get("file_name")
            lines.extend([f"## File: {current_file_name}", ""])
        lines.extend([f"### Page {page.get('page_no', '')}", ""])
        if options.get("final_text_only"):
            lines.extend(["```text", _text_for_page(page), "```", ""])
            continue
        if options.get("include_ocr_text"):
            lines.extend(["#### OCR 原始文本", "", "```text", page.get("ocr_text", ""), "```", ""])
        if options.get("include_corrected_text"):
            lines.extend(["#### LLM 校对文本", "", "```text", page.get("corrected_text", ""), "```", ""])
        if options.get("include_final_text"):
            lines.extend(["#### 用户最终确认文本", "", "```text", page.get("final_text", ""), "```", ""])
        if options.get("include_suggestions"):
            lines.extend(["#### 修订建议", "", "| 行号 | 原文本 | 建议文本 | 类型 | 置信度 | 状态 | 理由 |", "| -- | --- | ---- | ---- | ---- | ---- | --- |"])
            for s in page.get("suggestions", []):
                lines.append("| {line} | {original} | {suggested} | {cat} | {conf} | {status} | {reason} |".format(
                    line=_md_escape_cell(s.get("line_no", "")),
                    original=_md_escape_cell(s.get("original", "")),
                    suggested=_md_escape_cell(s.get("suggested", "")),
                    cat=_md_escape_cell(s.get("category", "")),
                    conf=_md_escape_cell(s.get("confidence", "")),
                    status=_md_escape_cell(s.get("status", "pending")),
                    reason=_md_escape_cell(s.get("reason", "")),
                ))
            lines.append("")
        if options.get("include_ocr_blocks"):
            lines.extend(["#### OCR 文本块", "", "| BlockID | Text | Confidence | BBox |", "| ------- | ---- | ---------- | ---- |"])
            for i, block in enumerate(page.get("ocr_blocks", []), start=1):
                lines.append(f"| {i} | {_md_escape_cell(block.get('text', ''))} | {_md_escape_cell(block.get('confidence', ''))} | {_md_escape_cell(_bbox_str(block))} |")
            lines.append("")
    export_path = Path(export_path)
    export_path.parent.mkdir(parents=True, exist_ok=True)
    export_path.write_text("\n".join(lines), encoding="utf-8")
    return str(export_path)


def export_to_docx(project_data: dict[str, Any], export_path: str | Path, scope: str = "project", options: dict | None = None) -> str:
    options = _opts(options)
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError("未安装 python-docx。请先运行：python -m pip install python-docx") from exc
    meta = _project_meta(project_data)
    doc = Document()
    doc.add_heading("BFSU ProofLens OCR 校对结果", level=0)
    doc.add_heading("一、基本信息", level=1)
    for key, label in [("project_name", "项目名称"), ("export_time", "导出时间"), ("ocr_backend", "OCR 引擎"), ("use_llm", "是否使用 LLM"), ("llm_model", "LLM 模型")]:
        doc.add_paragraph(f"{label}：{meta[key]}")
    for file_entry, _, page in _iter_pages(project_data, scope, options):
        doc.add_heading(f"{file_entry.get('file_name', '')} - Page {page.get('page_no', '')}", level=1)
        if options.get("include_ocr_text"):
            doc.add_heading("二、OCR 原始文本", level=2)
            doc.add_paragraph(page.get("ocr_text", ""))
        if options.get("include_corrected_text"):
            doc.add_heading("三、LLM 校对文本", level=2)
            doc.add_paragraph(page.get("corrected_text", ""))
        if options.get("include_final_text"):
            doc.add_heading("四、用户最终确认文本", level=2)
            doc.add_paragraph(page.get("final_text", ""))
        if options.get("include_suggestions"):
            doc.add_heading("五、修订建议", level=2)
            table = doc.add_table(rows=1, cols=6)
            hdr = table.rows[0].cells
            for i, h in enumerate(["行号", "原文本", "建议文本", "类型", "置信度", "理由"]):
                hdr[i].text = h
            for s in page.get("suggestions", []):
                row = table.add_row().cells
                row[0].text = str(s.get("line_no", ""))
                row[1].text = str(s.get("original", ""))
                row[2].text = str(s.get("suggested", ""))
                row[3].text = str(s.get("category", ""))
                row[4].text = str(s.get("confidence", ""))
                row[5].text = str(s.get("reason", ""))
    export_path = Path(export_path)
    export_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(export_path))
    return str(export_path)


def export_to_xlsx(project_data: dict[str, Any], export_path: str | Path, scope: str = "project", options: dict | None = None) -> str:
    options = _opts(options)
    try:
        from openpyxl import Workbook  # type: ignore
    except Exception as exc:
        raise RuntimeError("未安装 openpyxl。请先运行：python -m pip install openpyxl") from exc
    meta = _project_meta(project_data)
    pages = list(_iter_pages(project_data, scope, options))
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    total_suggestions = sum(len(p.get("suggestions", [])) for _, _, p in pages)
    accepted = sum(sum(1 for s in p.get("suggestions", []) if s.get("status") == "accepted") for _, _, p in pages)
    summary_rows = [
        ["ProjectName", meta["project_name"]],
        ["ExportTime", meta["export_time"]],
        ["FileCount", len({f.get("id") for f, _, _ in pages})],
        ["PageCount", len(pages)],
        ["OCRBackend", meta["ocr_backend"]],
        ["UseLLM", meta["use_llm"]],
        ["LLMModel", meta["llm_model"]],
        ["TotalSuggestions", total_suggestions],
        ["AcceptedSuggestions", accepted],
    ]
    for row in summary_rows:
        ws.append(row)

    wp = wb.create_sheet("Pages")
    wp.append(["FileName", "FilePath", "PageNo", "OCRBackend", "UseLLM", "LLMModel", "OCRStatus", "ProofreadStatus", "OCRText", "CorrectedText", "FinalText", "OCRTime", "LLMTime", "SuggestionCount", "AcceptedSuggestionCount", "CharacterCount", "Timestamp"])
    for file_entry, _, page in pages:
        suggs = page.get("suggestions", [])
        wp.append([
            file_entry.get("file_name", ""), file_entry.get("file_path", ""), page.get("page_no", ""),
            page.get("ocr_backend", meta["ocr_backend"]), meta["use_llm"], page.get("llm_model", meta["llm_model"]),
            page.get("ocr_status", ""), page.get("proofread_status", ""), page.get("ocr_text", ""),
            page.get("corrected_text", ""), page.get("final_text", ""), page.get("ocr_time", 0), page.get("llm_time", 0),
            len(suggs), sum(1 for s in suggs if s.get("status") == "accepted"), len(_text_for_page(page)), page.get("timestamp", "")
        ])

    wsugg = wb.create_sheet("Suggestions")
    wsugg.append(["FileName", "PageNo", "LineNo", "Original", "Suggested", "Reason", "Category", "Confidence", "Status", "Timestamp"])
    for file_entry, _, page in pages:
        for s in page.get("suggestions", []):
            wsugg.append([file_entry.get("file_name", ""), page.get("page_no", ""), s.get("line_no", ""), s.get("original", ""), s.get("suggested", ""), s.get("reason", ""), s.get("category", ""), s.get("confidence", ""), s.get("status", "pending"), s.get("timestamp", "")])

    wblo = wb.create_sheet("OCRBlocks")
    wblo.append(["FileName", "PageNo", "BlockID", "Text", "Confidence", "X1", "Y1", "X2", "Y2", "X3", "Y3", "X4", "Y4", "ReadingOrder"])
    for file_entry, _, page in pages:
        for i, block in enumerate(page.get("ocr_blocks", []), start=1):
            bbox = block.get("bbox") or []
            coords = []
            for j in range(4):
                point = bbox[j] if j < len(bbox) else ["", ""]
                coords.extend([point[0] if len(point) > 0 else "", point[1] if len(point) > 1 else ""])
            wblo.append([file_entry.get("file_name", ""), page.get("page_no", ""), i, block.get("text", ""), block.get("confidence", ""), *coords, block.get("reading_order", i)])

    export_path = Path(export_path)
    export_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(export_path))
    return str(export_path)
