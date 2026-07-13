# -*- coding: utf-8 -*-
"""Export utilities for BFSU WebLens."""
from __future__ import annotations

import csv
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Iterable

FIELDS = [
    "collected_at", "search_engine", "query", "query_raw", "search_vertical",
    "source_filter", "sort_mode", "site_limit", "date_filter_type", "date_start",
    "date_end", "start_ts", "end_ts", "baidu_gpc", "shard_start", "shard_end",
    "page", "rank", "title", "source", "actual_domain", "published_time",
    "link", "snippet", "search_url", "language_lr", "country_cr",
    "content_status", "content_word_count", "content_quality_score", "content_error",
    "content_extraction_method", "content_cleaning_scheme", "raw_html_path",
    "raw_text_path", "clean_text_path", "metadata_path", "metadata_excel_path"
]

def record_dicts(records: Iterable) -> list[dict]:
    rows = []
    extra_fields = [
        "content_status", "content_word_count", "content_quality_score", "content_error",
        "content_extraction_method", "content_cleaning_scheme", "raw_html_path",
        "raw_text_path", "clean_text_path", "metadata_path", "metadata_excel_path",
    ]
    for rec in records:
        if hasattr(rec, "to_dict"):
            row = rec.to_dict()
            for field in extra_fields:
                if hasattr(rec, field):
                    row[field] = getattr(rec, field)
            rows.append(row)
        elif isinstance(rec, dict):
            rows.append(rec)
    return rows

def export_records(records: Iterable, output_path: str, fmt: str) -> None:
    rows = record_dicts(records)
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    fmt = fmt.lower().strip().lstrip(".")
    if fmt == "xlsx":
        export_xlsx(rows, path)
    elif fmt == "csv":
        export_csv(rows, path)
    elif fmt == "txt":
        export_txt(rows, path)
    elif fmt == "docx":
        export_docx(rows, path)
    elif fmt == "xml":
        export_xml(rows, path)
    else:
        raise ValueError(f"Unsupported format: {fmt}")

def export_xlsx(rows: list[dict], path: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    wb = Workbook()
    ws = wb.active
    ws.title = "WebLens Results"
    ws.append(FIELDS)
    for row in rows:
        ws.append([row.get(f, "") for f in FIELDS])
    fill = PatternFill("solid", fgColor="17384A")
    font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center")
    for idx, field in enumerate(FIELDS, start=1):
        col = get_column_letter(idx)
        if field in {"title", "link", "snippet", "search_url"}:
            width = 60
        elif field.endswith("path") or field in {"baidu_gpc", "query_raw"}:
            width = 45
        elif field in {"collected_at", "published_time", "actual_domain"}:
            width = 22
        else:
            width = 16
        ws.column_dimensions[col].width = width
    ws.freeze_panes = "A2"
    wb.save(path)

def export_csv(rows: list[dict], path: Path) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=FIELDS)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in FIELDS})

def export_txt(rows: list[dict], path: Path) -> None:
    with path.open("w", encoding="utf-8") as f:
        for i, row in enumerate(rows, 1):
            f.write(f"[{i}] {row.get('title','')}\n")
            f.write(f"Time: {row.get('published_time','')} | Source: {row.get('source','')}\n")
            f.write(f"URL: {row.get('link','')}\n")
            f.write(f"Collected: {row.get('collected_at','')} | Shard: {row.get('shard_start','')}~{row.get('shard_end','')} | Page: {row.get('page','')} | Rank: {row.get('rank','')}\n")
            if row.get('snippet'):
                f.write(f"Snippet: {row.get('snippet','')}\n")
            f.write("\n")

def export_docx(rows: list[dict], path: Path) -> None:
    from docx import Document
    doc = Document()
    doc.add_heading("BFSU WebLens Results", level=1)
    for i, row in enumerate(rows, 1):
        doc.add_heading(f"{i}. {row.get('title','')}", level=2)
        p = doc.add_paragraph()
        p.add_run("Time: ").bold = True; p.add_run(str(row.get('published_time','')))
        p = doc.add_paragraph()
        p.add_run("Source: ").bold = True; p.add_run(str(row.get('source','')))
        p = doc.add_paragraph()
        p.add_run("URL: ").bold = True; p.add_run(str(row.get('link','')))
        p = doc.add_paragraph()
        p.add_run("Collected: ").bold = True; p.add_run(str(row.get('collected_at','')))
        if row.get('snippet'):
            doc.add_paragraph(str(row.get('snippet','')))
    doc.save(path)

def export_xml(rows: list[dict], path: Path) -> None:
    root = ET.Element("weblens_results")
    for row in rows:
        item = ET.SubElement(root, "record")
        for field in FIELDS:
            child = ET.SubElement(item, field)
            child.text = str(row.get(field, ""))
    tree = ET.ElementTree(root)
    tree.write(path, encoding="utf-8", xml_declaration=True)
