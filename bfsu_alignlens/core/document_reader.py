from __future__ import annotations

import re
from pathlib import Path
from typing import Tuple
from .utils import strip_markdown

SUPPORTED_EXTENSIONS = {'.txt', '.md', '.docx', '.doc', '.rtf'}


UTF8_NO_BOM = "utf-8"


def detect_encoding(path: str) -> str:
    """Return the only supported text import encoding.

    AlignLens intentionally reads plain text-like inputs as UTF-8 without BOM.
    This avoids silent GBK/Big5/UTF-16 guessing and keeps imported corpus text
    reproducible across machines. A UTF-8 BOM, if present in the file content,
    is removed later by normalize_text().
    """
    return UTF8_NO_BOM


def read_txt(path: str) -> str:
    try:
        return Path(path).read_text(encoding=UTF8_NO_BOM)
    except UnicodeDecodeError as exc:
        raise RuntimeError(
            f"Text import uses UTF-8 without BOM only. Please convert this file to UTF-8 and import again: {path}"
        ) from exc


def read_docx(path: str) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError('python-docx is not installed. Please run: pip install python-docx') from exc
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    # Include table cells when present.
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append('\t'.join(cell.text for cell in row.cells))
    return '\n'.join(paragraphs)


def read_rtf(path: str) -> str:
    raw = read_txt(path)
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
        return rtf_to_text(raw)
    except Exception:
        text = re.sub(r'{\\\*?[^{}]+}|[{}]', '', raw)
        text = re.sub(r'\\[a-zA-Z]+-?\d* ?', '', text)
        text = text.replace('\\par', '\n')
        return text


def read_doc(path: str) -> str:
    # Legacy .doc is deliberately handled gracefully. On Windows, users may install pywin32/LibreOffice.
    raise RuntimeError('Legacy .doc reading is not available in the current environment. Please convert this file to .docx and import again.')


def normalize_text(text: str, remove_excessive_spaces: bool = True, preserve_paragraphs: bool = True) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n').replace('\ufeff', '')
    if remove_excessive_spaces:
        lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.split('\n')]
        text = '\n'.join(lines)
    if preserve_paragraphs:
        text = re.sub(r'\n{3,}', '\n\n', text)
    else:
        text = re.sub(r'\n+', '\n', text)
    return text.strip()


def read_document(path: str, *, remove_excessive_spaces: bool = True, preserve_paragraphs: bool = True) -> str:
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise RuntimeError(f'Unsupported file format: {ext}')
    if ext == '.txt':
        text = read_txt(path)
    elif ext == '.md':
        text = strip_markdown(read_txt(path))
    elif ext == '.docx':
        text = read_docx(path)
    elif ext == '.rtf':
        text = read_rtf(path)
    elif ext == '.doc':
        text = read_doc(path)
    else:
        text = read_txt(path)
    return normalize_text(text, remove_excessive_spaces=remove_excessive_spaces, preserve_paragraphs=preserve_paragraphs)


def text_stats(text: str) -> Tuple[int, int]:
    chars = len(text)
    paras = len([p for p in re.split(r'\n\s*\n|\n+', text) if p.strip()])
    return chars, paras
