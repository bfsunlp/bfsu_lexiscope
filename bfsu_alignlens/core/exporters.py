from __future__ import annotations

import json
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Iterable, List

from .datatypes import AlignmentUnit
from .language_registry import tmx_code
from .multi_txt_exporter import MultiTxtOptions, export_multi_txt


def all_languages(units: List[AlignmentUnit]) -> List[str]:
    langs = []
    for u in units:
        for lang in u.segments.keys():
            if lang not in langs:
                langs.append(lang)
    return langs


def position_summary(unit: AlignmentUnit) -> str:
    parts = []
    for col, vals in (getattr(unit, 'positions', {}) or {}).items():
        if vals:
            parts.append(f"{col}:{','.join(str(v) for v in vals)}")
    if parts:
        return ' | '.join(parts)
    src = ','.join(str(x) for x in (unit.source_ids or []))
    tgts = ';'.join(f"{k}:{','.join(str(x) for x in v)}" for k, v in (unit.target_ids or {}).items())
    return f"source:{src} | {tgts}" if (src or tgts) else ''


def export_single_txt(units: List[AlignmentUnit], path: str):
    langs = all_languages(units)
    with open(path, 'w', encoding='utf-8') as f:
        for u in units:
            f.write(f'序号：{u.row_id}\n')
            f.write(f'对齐组：{u.group_id}\n')
            for lang in langs:
                f.write(f'{lang}：{u.segments.get(lang, "").strip()}\n')
            f.write(f'相似度：{u.similarity:.4f}\n')
            f.write(f'状态：{u.status}\n')
            if u.note:
                f.write(f'备注：{u.note}\n')
            f.write('\n')


def export_line_txt(units: List[AlignmentUnit], path: str, include_position_info: bool = False):
    langs = all_languages(units)
    with open(path, 'w', encoding='utf-8') as f:
        header = ['group_id', 'row_id'] + langs + ['similarity', 'status', 'note']
        f.write('\t'.join(header) + '\n')
        for u in units:
            row = [u.group_id, str(u.row_id)] + [u.segments.get(lang, '') for lang in langs] + [f'{u.similarity:.4f}', u.status, u.note]
            f.write('\t'.join(row) + '\n')


def export_json(units: List[AlignmentUnit], path: str):
    Path(path).write_text(json.dumps([u.to_dict() for u in units], ensure_ascii=False, indent=2), encoding='utf-8')


def export_xml(units: List[AlignmentUnit], path: str, include_position_info: bool = False):
    root = ET.Element('alignment_project')
    for u in units:
        au = ET.SubElement(root, 'alignment_unit', id=str(u.row_id), group_id=u.group_id, similarity=f'{u.similarity:.4f}', status=u.status, confirmed=str(u.confirmed).lower(), level=getattr(u, 'alignment_level', 'sentence'))
        for lang, text in u.segments.items():
            seg = ET.SubElement(au, 'seg', lang=lang)
            seg.text = text
        if u.note:
            note = ET.SubElement(au, 'note')
            note.text = u.note
    ET.ElementTree(root).write(path, encoding='utf-8', xml_declaration=True)


def export_tmx(units: List[AlignmentUnit], path: str, include_position_info: bool = False):
    root = ET.Element('tmx', version='1.4b')
    header = ET.SubElement(root, 'header', creationtool='BFSU AlignLens', creationtoolversion='1.1.0', datatype='PlainText', segtype='sentence', adminlang='en', srclang='mul', o_tmf='BFSUAlignLens')
    body = ET.SubElement(root, 'body')
    for u in units:
        tu = ET.SubElement(body, 'tu', tuid=f'{u.group_id}_{u.row_id}', usagecount='1')
        ET.SubElement(tu, 'prop', type='similarity').text = f'{u.similarity:.4f}'
        ET.SubElement(tu, 'prop', type='status').text = u.status
        ET.SubElement(tu, 'prop', type='alignment_level').text = getattr(u, 'alignment_level', 'sentence')
        for lang, text in u.segments.items():
            tuv = ET.SubElement(tu, 'tuv', {'xml:lang': tmx_code(lang)})
            ET.SubElement(tuv, 'seg').text = text
    ET.ElementTree(root).write(path, encoding='utf-8', xml_declaration=True)


def export_excel(units: List[AlignmentUnit], path: str, include_position_info: bool = False):
    try:
        from openpyxl import Workbook  # type: ignore
    except Exception as exc:
        raise RuntimeError('openpyxl is not installed. Please run: pip install openpyxl') from exc
    langs = all_languages(units)
    wb = Workbook()
    ws = wb.active
    ws.title = 'alignments'
    header = ['group_id', 'row_id', 'alignment_level', 'source_lang'] + langs + ['similarity', 'status', 'issue_type', 'note', 'llm_suggestion', 'confirmed']
    ws.append(header)
    for u in units:
        ws.append([u.group_id, u.row_id, getattr(u, 'alignment_level', 'sentence'), u.source_lang] + [u.segments.get(lang, '') for lang in langs] + [u.similarity, u.status, u.issue_type, u.note, u.llm_suggestion, u.confirmed])
    for col in ws.columns:
        try:
            ws.column_dimensions[col[0].column_letter].width = min(max(12, max(len(str(c.value or '')) for c in col[:100]) + 2), 60)
        except Exception:
            pass
    wb.save(path)


def export_word(units: List[AlignmentUnit], path: str, include_position_info: bool = False):
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError('python-docx is not installed. Please run: pip install python-docx') from exc
    langs = all_languages(units)
    doc = Document()
    doc.add_heading('BFSU AlignLens Alignment Results', level=1)
    table = doc.add_table(rows=1, cols=2 + len(langs) + 2)
    hdr = table.rows[0].cells
    cols = ['Group', 'Row'] + langs + ['Similarity', 'Status']
    for i, name in enumerate(cols):
        hdr[i].text = name
    for u in units:
        cells = table.add_row().cells
        vals = [u.group_id, str(u.row_id)] + [u.segments.get(lang, '') for lang in langs] + [f'{u.similarity:.4f}', u.status]
        for i, v in enumerate(vals):
            cells[i].text = v
    doc.save(path)


def export_by_extension(units: List[AlignmentUnit], path: str):
    ext = Path(path).suffix.lower()
    if ext == '.xlsx':
        return export_excel(units, path)
    if ext == '.txt':
        return export_line_txt(units, path)
    if ext == '.tmx':
        return export_tmx(units, path)
    if ext == '.xml':
        return export_xml(units, path)
    if ext == '.docx':
        return export_word(units, path)
    if ext == '.json':
        return export_json(units, path)
    return export_line_txt(units, path)
